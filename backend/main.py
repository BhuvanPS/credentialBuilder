"""
Main FastAPI server for the Credential Builder backend.

Handles:
1. Local file upload caching.
2. LinkedIn profile scraping & extraction.
3. Azure Blob Storage integration (with SAS token generation).
4. Document analysis extraction using Azure Content Understanding.
5. AI Agent executive summary synthesis using Azure AI Projects (or local fallback).
"""

import os
import re
import sys
import uuid
import json
from datetime import datetime, timedelta
from pathlib import Path
import httpx
from dotenv import load_dotenv

# Ensure the backend directory is in the sys.path so modules like 'db' are always importable
backend_dir = str(Path(__file__).resolve().parent)
if backend_dir not in sys.path:
    sys.path.insert(0, backend_dir)

# Load repository-root .env config values before importing other modules
env_path = Path(__file__).resolve().parents[1] / ".env"
load_dotenv(dotenv_path=env_path)

from azure.core.exceptions import ResourceExistsError
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Any, Dict, List, Optional
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
from openai import OpenAI
from azure.ai.projects import AIProjectClient
from azure.ai.inference import ChatCompletionsClient
from azure.ai.inference.models import SystemMessage, UserMessage, JsonSchemaFormat
from azure.core.credentials import AzureKeyCredential
from pydantic import ValidationError
import io
from db import init_db, save_candidate, list_candidates, get_candidate, delete_candidate
import docx
from fpdf import FPDF
from pypdf import PdfWriter

from cu_backend.profileAnalyser import analyze_url



app = FastAPI(title="Credential Builder Backend")

# Allow cross-origin requests from the React dev server ports
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

LINKEDIN_EXTRACTOR_URL = os.environ.get("LINKEDIN_EXTRACTOR_URL", "http://127.0.0.1:8000")
AZURE_STORAGE_CONNECTION_STRING = os.environ.get("AZURE_STORAGE_CONNECTION_STRING", "")
AZURE_STORAGE_ACCOUNT_NAME = os.environ.get("AZURE_STORAGE_ACCOUNT_NAME", "")
AZURE_STORAGE_ACCOUNT_KEY = os.environ.get("AZURE_STORAGE_ACCOUNT_KEY", "")
AZURE_STORAGE_CONTAINER_NAME = os.environ.get("AZURE_STORAGE_CONTAINER_NAME", "uploaded-files")

# In-memory document storage (temporary cache before saving to Azure)
uploaded_files: Dict[str, bytes] = {}
uploaded_file_names: Dict[str, str] = {}


# --- PYDANTIC SCHEMAS ---

class LinkedInRequest(BaseModel):
    url: str


class AzureUploadRequest(BaseModel):
    file_id: str
    file_name: str


class AnalyzeRequest(BaseModel):
    url: str


class GenerateSummaryRequest(BaseModel):
    name: str
    title: str
    summary: str
    coreCompetencies: List[str]
    keyExpertise: List[str]


class CredentialSummaryResponse(BaseModel, extra="forbid"):
    """Structured output schema for the AI-synthesized executive summary."""
    name: str
    title: str
    summary: str


class CandidateSaveRequest(BaseModel):
    name: str
    title: str
    profile_picture_url: Optional[str] = None
    form_data: Dict[str, Any]
    summary_data: Dict[str, Any]


# --- HELPER UTILITIES ---

def parse_connection_string(connection_string: str) -> Dict[str, str]:
    """Parses key-value pairs from an Azure Storage connection string."""
    parts = [segment.strip() for segment in connection_string.split(";") if segment.strip()]
    values = {}
    for segment in parts:
        if "=" not in segment:
            continue
        key, value = segment.split("=", 1)
        values[key.strip()] = value.strip()
    return values


def get_blob_service_client() -> BlobServiceClient:
    """Initializes and returns the Azure BlobServiceClient."""
    if AZURE_STORAGE_CONNECTION_STRING:
        return BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)

    if not AZURE_STORAGE_ACCOUNT_NAME or not AZURE_STORAGE_ACCOUNT_KEY:
        raise HTTPException(
            status_code=500,
            detail="Azure storage credentials are not configured. Set AZURE_STORAGE_CONNECTION_STRING or credentials.",
        )

    account_url = f"https://{AZURE_STORAGE_ACCOUNT_NAME}.blob.core.windows.net"
    return BlobServiceClient(account_url=account_url, credential=AZURE_STORAGE_ACCOUNT_KEY)


def get_sas_url(blob_name: str) -> str:
    """Generates a SAS (Shared Access Signature) URL for a given blob name (expires in 1 hour)."""
    expiry_hours = int(os.getenv("AZURE_SAS_EXPIRY_HOURS", "1"))

    account_name = AZURE_STORAGE_ACCOUNT_NAME
    account_key = AZURE_STORAGE_ACCOUNT_KEY

    if AZURE_STORAGE_CONNECTION_STRING:
        parsed = parse_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        account_name = parsed.get("AccountName")
        account_key = parsed.get("AccountKey")

    if not account_name or not account_key:
        raise HTTPException(
            status_code=500,
            detail="Azure storage credentials are not configured.",
        )

    sas_token = generate_blob_sas(
        account_name=account_name,
        container_name=AZURE_STORAGE_CONTAINER_NAME,
        blob_name=blob_name,
        account_key=account_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(hours=expiry_hours),
    )
    return f"https://{account_name}.blob.core.windows.net/{AZURE_STORAGE_CONTAINER_NAME}/{blob_name}?{sas_token}"


def clean_json_text(text: str) -> str:
    """Removes markdown code fences (e.g. ```json) around a JSON output string."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    return text


# --- API ROUTES ---

def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Converts DOCX file bytes into PDF bytes using python-docx and fpdf2."""
    doc = docx.Document(io.BytesIO(docx_bytes))
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)
    
    # Write paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Clean text to latin-1 to avoid encoding crashes
            safe_text = text.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(0, 5, safe_text)
            pdf.ln(3)
            
    # Write tables
    for table in doc.tables:
        pdf.ln(5)
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                safe_text = row_text.encode("latin-1", "replace").decode("latin-1")
                pdf.multi_cell(0, 5, safe_text)
                pdf.ln(3)
                
    return bytes(pdf.output())


def merge_pdfs(pdf_list: List[bytes]) -> bytes:
    """Merges multiple PDF bytearrays/bytes into a single PDF bytearray using pypdf PdfWriter."""
    writer = PdfWriter()
    for pdf_bytes in pdf_list:
        writer.append(io.BytesIO(pdf_bytes))
    
    out = io.BytesIO()
    writer.write(out)
    writer.close()
    return out.getvalue()


def generate_meaningful_filename(original_names: List[str], linkedin_url: str = None) -> str:
    """
    Generates a clean, human-readable filename for merged/single documents.

    Priority order:
      1. LinkedIn URL slug  (stripped of trailing LinkedIn hash IDs)
      2. Uploaded resume/CV filename stem
      3. Any uploaded filename stem
      4. Generic fallback
    """
    name_slug = ""

    # --- 1. Extract slug from LinkedIn URL ---
    if linkedin_url:
        match = re.search(r"/in/([^/?#]+)", linkedin_url)
        if match:
            raw_slug = match.group(1).strip()
            # LinkedIn sometimes appends a short alphanumeric hash like "-a1b2c3d4"
            # Pattern: trailing segment that is 7-10 hex/alnum chars after the last hyphen
            raw_slug = re.sub(r'-[a-f0-9]{6,10}$', '', raw_slug, flags=re.IGNORECASE)
            # Also strip query string fragments that may have leaked
            raw_slug = raw_slug.split('?')[0].split('#')[0].strip()
            if raw_slug:
                name_slug = raw_slug

    # --- 2. Fall back to uploaded file name stem ---
    if not name_slug and original_names:
        # Prefer files whose name contains resume/CV keywords
        keywords = ['resume', 'cv', 'profile', 'bio', 'background', 'credentials']
        primary_name = None
        for filename in original_names:
            lower_name = os.path.splitext(filename)[0].lower()
            if any(kw in lower_name for kw in keywords):
                primary_name = filename
                break
        if not primary_name:
            primary_name = original_names[0]
        name_slug = os.path.splitext(primary_name)[0]

    # --- 3. Ultimate fallback ---
    if not name_slug:
        name_slug = "candidate_profile"

    # --- Normalise the slug ---
    # Replace hyphens with underscores for consistency
    name_slug = name_slug.replace('-', '_')
    # Keep only alphanumeric, underscores, and spaces; strip everything else
    clean_slug = "".join(c for c in name_slug if c.isalnum() or c in ('_', ' ')).strip()
    # Collapse multiple underscores/spaces into a single underscore
    clean_slug = re.sub(r'[_ ]+', '_', clean_slug).strip('_')
    # Truncate to avoid excessively long blob names
    clean_slug = clean_slug[:60]

    if not clean_slug:
        clean_slug = "candidate_profile"

    return f"{clean_slug}_profile.pdf"


@app.post("/upload")
async def upload_documents(
    files: List[UploadFile] = File(None),
    linkedin_file_id: str = Form(None),
    linkedin_url: str = Form(None)
) -> Dict[str, Any]:
    """
    Accepts multiple documents (PDF, DOCX) and/or a previously fetched LinkedIn PDF ID.
    Converts DOCX to PDF, merges all files using pypdf, and caches the result.
    """
    processed_pdfs = []
    original_names = []

    # 1. Retrieve LinkedIn PDF if file ID is provided
    if linkedin_file_id:
        if linkedin_file_id in uploaded_files:
            processed_pdfs.append(uploaded_files[linkedin_file_id])
            # Use the actual cached filename (e.g. "john-doe.pdf") not a hardcoded generic name
            cached_linkedin_name = uploaded_file_names.get(linkedin_file_id, "linkedin_profile.pdf")
            original_names.append(cached_linkedin_name)
        else:
            raise HTTPException(
                status_code=400,
                detail="LinkedIn file ID not found in cache"
            )

    # 2. Process uploaded files
    if files:
        for file in files:
            if not file.filename:
                continue
                
            contents = await file.read()
            if len(contents) > 10 * 1024 * 1024:
                raise HTTPException(
                    status_code=400, 
                    detail=f"File {file.filename} exceeds the 10MB limit."
                )
                
            ext = file.filename.split('.')[-1].lower()
            if ext == 'docx':
                try:
                    pdf_bytes = convert_docx_to_pdf(contents)
                    processed_pdfs.append(pdf_bytes)
                    original_names.append(file.filename)
                except Exception as e:
                    raise HTTPException(
                        status_code=500, 
                        detail=f"Failed to convert {file.filename} to PDF: {str(e)}"
                    )
            elif ext == 'pdf':
                processed_pdfs.append(contents)
                original_names.append(file.filename)
            else:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Unsupported format for file {file.filename}. Only PDF and DOCX are allowed."
                )

    if not processed_pdfs:
        raise HTTPException(status_code=400, detail="No valid documents or profile data provided")

    # 3. Merge or return single file
    if len(processed_pdfs) == 1:
        final_pdf = processed_pdfs[0]
        base_name = original_names[0]
        # Convert docx stem to pdf extension
        stem, ext = os.path.splitext(base_name)
        if ext.lower() == '.docx':
            filename = stem + '.pdf'
        elif ext.lower() == '.pdf':
            # If this came from LinkedIn scraper the name is already meaningful (e.g. john-doe.pdf)
            # Run it through the normaliser so it is consistent with the merged-file path
            filename = generate_meaningful_filename([base_name], linkedin_url)
        else:
            filename = generate_meaningful_filename([base_name], linkedin_url)
    else:
        try:
            final_pdf = merge_pdfs(processed_pdfs)
            filename = generate_meaningful_filename(original_names, linkedin_url)
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to merge PDF files: {str(e)}"
            )

    file_id = uuid.uuid4().hex
    uploaded_files[file_id] = final_pdf
    uploaded_file_names[file_id] = filename

    return {"file_id": file_id, "filename": filename}


@app.post("/download")
async def download_linkedin_profile(request: LinkedInRequest) -> Dict[str, Any]:
    """Triggers the LinkedIn extractor scraper, downloads the profile PDF, and caches it in-memory."""
    if not request.url:
        raise HTTPException(status_code=400, detail="LinkedIn URL is required")

    async with httpx.AsyncClient(timeout=120.0) as client:
        response = await client.post(
            f"{LINKEDIN_EXTRACTOR_URL}/download",
            json={"url": request.url},
        )

    if response.status_code != 200:
        raise HTTPException(status_code=response.status_code, detail=response.text)

    data = response.json()
    file_path = data.get("file_path")
    if not file_path:
        raise HTTPException(status_code=502, detail="Extractor returned no file path")

    filename = os.path.basename(file_path)
    file_id = uuid.uuid4().hex
    with open(file_path, "rb") as source_file:
        contents = source_file.read()

    uploaded_files[file_id] = contents
    uploaded_file_names[file_id] = filename

    return {"file_id": file_id, "file_path": file_path, "filename": filename}


@app.post("/azure/upload")
async def upload_to_azure(request: AzureUploadRequest) -> Dict[str, Any]:
    """Uploads cached file content directly to Azure Blob Storage and returns a SAS URL."""
    if request.file_id not in uploaded_files:
        raise HTTPException(status_code=400, detail="Invalid file_id")

    blob_service_client = get_blob_service_client()
    container_client = blob_service_client.get_container_client(AZURE_STORAGE_CONTAINER_NAME)
    try:
        container_client.create_container()
    except ResourceExistsError:
        pass

    blob_name = os.path.basename(request.file_name) if request.file_name else uploaded_file_names.get(request.file_id, "upload")
    blob_client = container_client.get_blob_client(blob_name)
    blob_client.upload_blob(uploaded_files[request.file_id], overwrite=True)

    sas_url = get_sas_url(blob_name)
    return {"blob_url": blob_client.url, "sas_url": sas_url}


@app.post("/analyze")
def analyze_blob(request: AnalyzeRequest) -> Dict[str, Any]:
    """Submits the Azure SAS URL to Azure Content Understanding to parse credentials schema."""
    if not request.url:
        raise HTTPException(status_code=400, detail="URL is required for analysis")

    try:
        result = analyze_url(request.url)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    return result


@app.post("/generate-summary")
async def generate_summary(request: GenerateSummaryRequest) -> Dict[str, Any]:
    """
    Synthesizes a high-impact executive summary using OpenAI client responses API.
    """
    try:
        endpoint = os.getenv(
            "AZURE_INFERENCE_ENDPOINT",
            "https://credentialbuilder-trial-resource.services.ai.azure.com/openai/v1",
        )
        deployment_name = os.getenv(
            "AZURE_DEPLOYMENT_NAME",
            "gpt-5.2-475108"
        )
        api_key = os.getenv("AZURE_INFERENCE_CREDENTIAL") or os.getenv("OPENAI_API_KEY")

        if api_key:
            client = OpenAI(
                base_url=endpoint,
                api_key=api_key
            )
        else:
            token_provider = get_bearer_token_provider(DefaultAzureCredential(), "https://ai.azure.com/.default")
            client = OpenAI(
                base_url=endpoint,
                api_key=token_provider
            )

        input_prompt = (
            f"You are an expert executive resume writer for a professional services firm. "
            f"Synthesize a compelling, concise professional summary for the candidate. "
            f"Return the candidate's name and title exactly as provided. "
            f"Write the summary in third-person, present tense, 1-3 sentences.\n\n"
            f"Candidate Info:\n"
            f"Name: {request.name}\n"
            f"Title: {request.title}\n"
            f"Raw Summary: {request.summary}\n"
            f"Core Competencies: {', '.join(request.coreCompetencies)}\n"
            f"Key Expertise: {', '.join(request.keyExpertise)}\n\n"
            f"You MUST format your response as a valid JSON object matching this schema:\n"
            f'{{"name": "string", "title": "string", "summary": "string"}}\n'
            f"Return only the JSON object."
        )

        response = client.responses.create(
            model=deployment_name,
            input=input_prompt,
        )

        output_item = response.output[0]
        raw_text = ""
        
        # Robustly extract content string from output content part (which may be a list of ContentParts)
        if hasattr(output_item, "content") and getattr(output_item, "content"):
            content_part = output_item.content
            if isinstance(content_part, list):
                parts = []
                for part in content_part:
                    if hasattr(part, "text") and getattr(part, "text"):
                        parts.append(part.text)
                    elif isinstance(part, dict) and "text" in part:
                        parts.append(part["text"])
                    elif isinstance(part, str):
                        parts.append(part)
                    else:
                        parts.append(str(part))
                raw_text = "".join(parts)
            else:
                raw_text = str(content_part)
        elif hasattr(output_item, "text") and getattr(output_item, "text"):
            raw_text = output_item.text
        elif isinstance(output_item, dict):
            raw_text = output_item.get("content") or output_item.get("text") or str(output_item)
        else:
            raw_text = str(output_item)

        raw_text = clean_json_text(raw_text)
        raw = json.loads(raw_text)

        return {
            "name": raw.get("name") or request.name,
            "title": raw.get("title") or request.title,
            "summary": raw.get("summary") or request.summary,
        }

    except Exception as exc:
        print(f"Azure OpenAI responses inference call failed ({exc}). Falling back to rule-based synthesis.")

    # ── Local rule-based fallback ────────────────────────────────────────────
    comps_str = ", ".join(request.coreCompetencies)
    exp_str = ", ".join(request.keyExpertise)

    fallback_summary = request.summary
    if exp_str:
        fallback_summary += f"\n\nSpecialized in: {exp_str}."
    if comps_str:
        fallback_summary += f"\nKey capabilities include: {comps_str}."

    return {
        "name": request.name,
        "title": request.title,
        "summary": fallback_summary,
    }


@app.on_event("startup")
def startup_event():
    init_db()


@app.get("/api/candidates")
def get_candidates():
    try:
        return list_candidates()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database query failed: {str(e)}")


@app.get("/api/candidates/{name}")
def get_candidate_by_name(name: str):
    try:
        candidate = get_candidate(name)
        if not candidate:
            raise HTTPException(status_code=404, detail="Candidate not found")
        return candidate
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database query failed: {str(e)}")


@app.post("/api/candidates")
def save_candidate_endpoint(req: CandidateSaveRequest):
    if not req.name.strip():
        raise HTTPException(status_code=400, detail="Candidate name is required")
    try:
        save_candidate(
            name=req.name.strip(),
            title=req.title.strip(),
            profile_picture_url=req.profile_picture_url,
            form_data=req.form_data,
            summary_data=req.summary_data
        )
        return {"status": "success", "message": "Candidate profile saved successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database write failed: {str(e)}")


@app.delete("/api/candidates/{name}")
def delete_candidate_endpoint(name: str):
    try:
        delete_candidate(name)
        return {"status": "success", "message": "Candidate deleted successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database deletion failed: {str(e)}")
